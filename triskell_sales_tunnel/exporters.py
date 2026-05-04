"""Exporteurs : .txt, .pdf, .docx.

Tous reçoivent les mêmes paramètres et écrivent un fichier sur disque.
DECISION: reportlab pour PDF (Unicode robuste, zéro asset). python-docx pour Word.
DECISION: la palette utilisée dans les exports reprend les couleurs Triskell pour
la cohérence visuelle (en-tête bleu marine, accent turquoise).
"""

from __future__ import annotations

import logging
from datetime import datetime
from pathlib import Path
from typing import Optional

LOG = logging.getLogger("triskell.exporters")

# Couleurs Triskell — dupliquées ici (sans importer theme.py qui pull tkinter).
NAVY_RGB = (10 / 255, 37 / 255, 64 / 255)
TURQUOISE_RGB = (0 / 255, 178 / 255, 169 / 255)
TEXT_RGB = (15 / 255, 23 / 255, 42 / 255)


# ---------------------------------------------------------------------------
# Modèle d'export
# ---------------------------------------------------------------------------

class ExportPayload:
    """Conteneur agnostique pour les 3 formats."""

    def __init__(
        self,
        *,
        body: str,
        subject: str,
        product_label: str,
        client_label: str,
        channel_label: str,
    ) -> None:
        self.body = body
        self.subject = subject
        self.product_label = product_label
        self.client_label = client_label
        self.channel_label = channel_label
        self.exported_at = datetime.now()

    @property
    def header_lines(self) -> list[str]:
        return [
            f"Produit : {self.product_label}",
            f"Cible   : {self.client_label}",
            f"Canal   : {self.channel_label}",
            f"Date    : {self.exported_at.strftime('%Y-%m-%d %H:%M')}",
        ]


# ---------------------------------------------------------------------------
# .txt
# ---------------------------------------------------------------------------

def export_txt(payload: ExportPayload, dest: Path) -> Path:
    lines = ["# Triskell Sales Tunnel — export"]
    lines += [f"# {l}" for l in payload.header_lines]
    lines.append("")
    if payload.subject:
        lines.append(f"Objet : {payload.subject}")
        lines.append("")
    lines.append(payload.body)
    dest.write_text("\n".join(lines) + "\n", encoding="utf-8")
    return dest


# ---------------------------------------------------------------------------
# .pdf  (reportlab)
# ---------------------------------------------------------------------------

def export_pdf(payload: ExportPayload, dest: Path) -> Path:
    from reportlab.lib.pagesizes import A4  # type: ignore[import-untyped]
    from reportlab.lib.units import mm  # type: ignore[import-untyped]
    from reportlab.pdfgen import canvas  # type: ignore[import-untyped]

    page_w, page_h = A4
    margin = 22 * mm

    c = canvas.Canvas(str(dest), pagesize=A4)
    c.setTitle(f"Triskell Sales Tunnel — {payload.channel_label}")
    c.setAuthor("Triskell Studio")

    # Bandeau d'en-tête bleu marine
    band_h = 26 * mm
    c.setFillColorRGB(*NAVY_RGB)
    c.rect(0, page_h - band_h, page_w, band_h, fill=1, stroke=0)

    # Titre branding
    c.setFillColorRGB(1, 1, 1)
    c.setFont("Helvetica-Bold", 17)
    c.drawString(margin, page_h - band_h + 14 * mm, "Triskell Studio")
    c.setFont("Helvetica", 11)
    c.drawString(margin, page_h - band_h + 7 * mm, "Sales Tunnel — Template de prospection")

    # Filet turquoise
    c.setFillColorRGB(*TURQUOISE_RGB)
    c.rect(0, page_h - band_h - 1.5 * mm, page_w, 1.5 * mm, fill=1, stroke=0)

    # Méta
    y = page_h - band_h - 12 * mm
    c.setFillColorRGB(*TEXT_RGB)
    c.setFont("Helvetica-Bold", 10)
    for line in payload.header_lines:
        c.drawString(margin, y, line)
        y -= 5 * mm

    # Objet
    y -= 4 * mm
    if payload.subject:
        c.setFont("Helvetica-Bold", 12)
        c.drawString(margin, y, "Objet")
        y -= 5 * mm
        c.setFont("Helvetica", 11)
        for line in _wrap(payload.subject, 90):
            c.drawString(margin, y, line)
            y -= 5 * mm
        y -= 3 * mm

    # Corps
    c.setFont("Helvetica-Bold", 12)
    c.drawString(margin, y, "Message")
    y -= 6 * mm

    c.setFont("Helvetica", 11)
    leading = 5 * mm
    line_w = 95
    for raw_line in payload.body.split("\n"):
        if not raw_line.strip():
            y -= leading
            continue
        for line in _wrap(raw_line, line_w):
            if y < margin + 10 * mm:
                _new_page(c, page_w, page_h, margin)
                y = page_h - margin
            c.drawString(margin, y, line)
            y -= leading

    # Pied de page
    c.setFillColorRGB(*TEXT_RGB)
    c.setFont("Helvetica-Oblique", 8)
    c.drawString(margin, 12 * mm, "Généré par Triskell Sales Tunnel · triskell-studio.fr")

    c.showPage()
    c.save()
    return dest


def _new_page(c, page_w: float, page_h: float, margin: float) -> None:  # type: ignore[no-untyped-def]
    from reportlab.lib.units import mm  # type: ignore[import-untyped]
    c.showPage()
    # Bandeau léger
    c.setFillColorRGB(*TURQUOISE_RGB)
    c.rect(0, page_h - 4 * mm, page_w, 4 * mm, fill=1, stroke=0)
    c.setFillColorRGB(*TEXT_RGB)
    c.setFont("Helvetica", 11)


def _wrap(text: str, width: int) -> list[str]:
    """Wrap simple par mots, sans dépendance externe."""
    words = text.split(" ")
    lines: list[str] = []
    cur = ""
    for w in words:
        if not cur:
            cur = w
        elif len(cur) + 1 + len(w) <= width:
            cur += " " + w
        else:
            lines.append(cur)
            cur = w
    if cur:
        lines.append(cur)
    return lines


# ---------------------------------------------------------------------------
# .docx  (python-docx)
# ---------------------------------------------------------------------------

def export_docx(payload: ExportPayload, dest: Path) -> Path:
    from docx import Document  # type: ignore[import-untyped]
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore[import-untyped]
    from docx.shared import Pt, RGBColor  # type: ignore[import-untyped]

    doc = Document()

    # Style global plus moderne
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Titre branding
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Triskell Studio")
    run.font.size = Pt(20)
    run.bold = True
    run.font.color.rgb = RGBColor(10, 37, 64)

    sub = doc.add_paragraph()
    run = sub.add_run("Sales Tunnel — Template de prospection")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0, 178, 169)

    # Méta
    meta = doc.add_paragraph()
    meta.add_run("\n".join(payload.header_lines)).font.size = Pt(9)

    # Objet
    if payload.subject:
        h = doc.add_paragraph()
        run = h.add_run("Objet")
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(10, 37, 64)
        doc.add_paragraph(payload.subject)

    # Corps
    h = doc.add_paragraph()
    run = h.add_run("Message")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(10, 37, 64)
    for chunk in payload.body.split("\n"):
        doc.add_paragraph(chunk)

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    fr = footer.add_run("Généré par Triskell Sales Tunnel · triskell-studio.fr")
    fr.italic = True
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor(120, 120, 120)

    doc.save(str(dest))
    return dest
